# -*- coding: utf-8 -*-
"""
生成Word分析报告
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import pandas as pd
import numpy as np

output_dir = "/Users/liaoweizhou/.openclaw/workspace/"

def set_chinese_font(run, font_name='SimSun', font_size=12, bold=False):
    """设置中文字体"""
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.font.bold = bold
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)

def add_heading_zh(doc, text, level=1):
    """添加中文标题"""
    heading = doc.add_heading(level=level)
    run = heading.add_run(text)
    font_names = ['SimHei', '黑体'] if level <= 2 else ['SimSun', '宋体']
    set_chinese_font(run, font_name=font_names[0], font_size=(18-level*2), bold=True)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return heading

def add_paragraph_zh(doc, text, bold=False, font_size=12, alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=0):
    """添加中文段落"""
    para = doc.add_paragraph()
    para.alignment = alignment
    para.paragraph_format.first_line_indent = Cm(first_line_indent)
    run = para.add_run(text)
    set_chinese_font(run, font_name='SimSun', font_size=font_size, bold=bold)
    return para

def add_table_zh(doc, data, headers=None):
    """添加中文表格"""
    if headers:
        table = doc.add_table(rows=1, cols=len(headers))
        hdr_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            hdr_cells[i].text = str(header)
            for paragraph in hdr_cells[i].paragraphs:
                for run in paragraph.runs:
                    set_chinese_font(run, font_name='SimHei', font_size=10, bold=True)
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        for row_data in data:
            row_cells = table.add_row().cells
            for i, value in enumerate(row_data):
                row_cells[i].text = str(value)
                for paragraph in row_cells[i].paragraphs:
                    for run in paragraph.runs:
                        set_chinese_font(run, font_name='SimSun', font_size=9)
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        table = doc.add_table(rows=len(data), cols=len(data[0]))
        for i, row_data in enumerate(data):
            for j, value in enumerate(row_data):
                table.cell(i, j).text = str(value)
    
    return table

def generate_report():
    """生成Word报告"""
    doc = Document()
    
    # 设置文档默认字体
    style = doc.styles['Normal']
    style.font.name = 'SimSun'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
    style.font.size = Pt(12)
    
    # ========== 封面 ==========
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('柳州市伊蚊布雷图监测分析报告')
    set_chinese_font(run, font_name='SimHei', font_size=22, bold=True)
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('(2021-2025年)')
    set_chinese_font(run, font_name='SimHei', font_size=16, bold=True)
    
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = date_para.add_run('报告日期：2025年2月')
    set_chinese_font(run, font_name='SimSun', font_size=12)
    
    doc.add_page_break()
    
    # ========== 一、数据概况 ==========
    add_heading_zh(doc, '一、数据概况', level=1)
    
    add_paragraph_zh(doc, 
        "本报告基于2021年至2025年柳州市病媒生物监测数据，对伊蚊布雷图指数（Breteau Index, BI）进行时空分布分析。布雷图指数是指每100户调查住户中发现的伊蚊阳性积水容器数，是评估登革热传播风险的重要指标。")
    
    add_paragraph_zh(doc, "", first_line_indent=0)
    
    add_heading_zh(doc, '1.1 监测基本情况', level=2)
    
    overview_data = [
        ["监测年份", "2021-2025年（共5年）"],
        ["监测区域", "柳州市各区县（含城中区、鱼峰区、柳南区、柳北区、柳江区、柳东新区、阳和新区、柳城县、鹿寨县、融安县、融水县、三江县等）"],
        ["监测频次", "每月监测"],
        ["累计调查户数", "74,606户"],
        ["累计发现阳性容器", "2,795个"],
        ["有效监测记录", "1,651条"],
        ["总体平均布雷图指数", "3.75"]
    ]
    
    for item in overview_data:
        p = add_paragraph_zh(doc, f"• {item[0]}：{item[1]}", first_line_indent=0)
    
    add_heading_zh(doc, '1.2 布雷图指数评价标准', level=2)
    
    add_paragraph_zh(doc, 
        "根据登革热防控相关技术指南，布雷图指数风险分级标准如下：", first_line_indent=0.74)
    
    # 风险标准表格
    standard_data = [
        ["风险等级", "布雷图指数标准", "防控建议"],
        ["安全", "BI < 5", "常规监测，维持现状"],
        ["风险", "5 ≤ BI < 10", "加强监测，清除积水容器"],
        ["高危", "BI ≥ 10", "立即采取控制措施，紧急灭蚊"]
    ]
    
    add_table_zh(doc, standard_data[1:], standard_data[0])
    
    doc.add_paragraph()
    
    # ========== 二、时间维度分析 ==========
    add_heading_zh(doc, '二、时间维度分析', level=1)
    
    add_heading_zh(doc, '2.1 年度变化趋势', level=2)
    
    add_paragraph_zh(doc, 
        "2021-2025年柳州市布雷图指数年度变化呈现明显上升趋势。2021-2022年布雷图指数处于较低水平（1.58-1.69），2023年开始上升至3.14，2024年略有回落但仍高于前期水平，2025年显著上升至5.34，首次超过风险阈值（BI=5），提示登革热传播风险显著增加。")
    
    add_paragraph_zh(doc, "", first_line_indent=0)
    
    # 年度统计表格
    yearly_data = [
        ["年份", "监测次数", "调查户数", "阳性容器", "平均BI", "最大值"],
        ["2021", "38", "6,713", "106", "1.58", "8.00"],
        ["2022", "41", "5,680", "96", "1.69", "9.29"],
        ["2023", "224", "6,778", "213", "3.14", "60.00"],
        ["2024", "410", "18,947", "431", "2.27", "200.00"],
        ["2025", "938", "36,488", "1,949", "5.34", "78.79"]
    ]
    add_table_zh(doc, yearly_data[1:], yearly_data[0])
    
    doc.add_paragraph()
    
    # 插入年度趋势图
    try:
        doc.add_picture(f'{output_dir}chart_01_yearly_trend.png', width=Inches(5.5))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_paragraph_zh(doc, "图1 2021-2025年布雷图指数年度变化趋势", font_size=10, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    except:
        add_paragraph_zh(doc, "[图表：年度趋势图]")
    
    add_heading_zh(doc, '2.2 月度分布特征', level=2)
    
    add_paragraph_zh(doc, 
        "布雷图指数呈现明显的季节性波动，夏季（6-8月）为高峰期。8月份平均布雷图指数最高（6.93），超过风险阈值；9月份次之（5.60）。4-5月和10-12月布雷图指数相对较低，处于安全水平。")
    
    add_paragraph_zh(doc, "", first_line_indent=0)
    
    # 月度统计表格
    monthly_data = [
        ["月份", "监测次数", "平均BI", "风险等级"],
        ["4月", "128", "1.60", "安全"],
        ["5月", "124", "2.98", "安全"],
        ["6月", "142", "3.93", "安全"],
        ["7月", "213", "4.05", "安全"],
        ["8月", "211", "6.93", "风险"],
        ["9月", "216", "5.60", "风险"],
        ["10月", "209", "3.76", "安全"],
        ["11月", "208", "2.33", "安全"],
        ["12月", "200", "0.81", "安全"]
    ]
    add_table_zh(doc, monthly_data[1:], monthly_data[0])
    
    doc.add_paragraph()
    
    # 插入月度分布图
    try:
        doc.add_picture(f'{output_dir}chart_02_monthly_distribution.png', width=Inches(5.5))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_paragraph_zh(doc, "图2 各月份布雷图指数分布特征", font_size=10, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    except:
        add_paragraph_zh(doc, "[图表：月度分布图]")
    
    add_heading_zh(doc, '2.3 季节性变化规律', level=2)
    
    add_paragraph_zh(doc, 
        "从季节性分布来看，夏季（6-8月）布雷图指数最高，平均为4.97，接近风险阈值；秋季（9-11月）次之，平均为3.94；春季（3-5月）平均为2.31；冬季（12-2月）最低，平均仅0.81。这一分布规律与伊蚊的活动季节高度吻合，高温多雨季节有利于蚊虫孳生。")
    
    # 季节统计表格
    season_data = [
        ["季节", "月份范围", "平均BI", "风险等级"],
        ["夏季", "6-8月", "4.97", "接近风险"],
        ["秋季", "9-11月", "3.94", "安全"],
        ["春季", "3-5月", "2.31", "安全"],
        ["冬季", "12-2月", "0.81", "安全"]
    ]
    add_table_zh(doc, season_data[1:], season_data[0])
    
    doc.add_page_break()
    
    # ========== 三、空间维度分析 ==========
    add_heading_zh(doc, '三、空间维度分析', level=1)
    
    add_heading_zh(doc, '3.1 区域对比分析', level=2)
    
    add_paragraph_zh(doc, 
        "各监测区域布雷图指数差异明显。柳城县平均布雷图指数最高（6.40），超过风险阈值；鹿寨县次之（5.69），同样处于风险水平。柳南区（4.82）、融安县（4.51）接近风险阈值。柳东新区（0.36）、阳和新区（0.82）等新建城区布雷图指数较低。")
    
    add_paragraph_zh(doc, "", first_line_indent=0)
    
    # 区域统计表格
    area_data = [
        ["监测区域", "监测次数", "平均BI", "风险等级"],
        ["柳城县", "117", "6.40", "风险"],
        ["鹿寨县", "250", "5.69", "风险"],
        ["柳南区", "144", "4.82", "接近风险"],
        ["融安县", "142", "4.51", "接近风险"],
        ["城中区", "145", "3.38", "安全"],
        ["鱼峰区", "142", "3.23", "安全"],
        ["三江县", "108", "3.08", "安全"],
        ["融水县", "290", "2.94", "安全"],
        ["柳北区", "144", "2.40", "安全"],
        ["柳江区", "123", "1.96", "安全"]
    ]
    add_table_zh(doc, area_data[1:], area_data[0])
    
    doc.add_paragraph()
    
    # 插入区域对比图
    try:
        doc.add_picture(f'{output_dir}chart_03_area_comparison.png', width=Inches(5.5))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_paragraph_zh(doc, "图3 各监测区域布雷图指数对比", font_size=10, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    except:
        add_paragraph_zh(doc, "[图表：区域对比图]")
    
    add_heading_zh(doc, '3.2 高风险区域识别', level=2)
    
    add_paragraph_zh(doc, 
        f"监测期间共发现449条高风险记录（BI≥5），占总监测记录的27.2%。其中高危记录（BI≥10）249条，占15.1%。高风险区域主要集中在柳城县、鹿寨县、融水县、柳南区、鱼峰区和融安县。")
    
    add_paragraph_zh(doc, "", first_line_indent=0)
    
    # 高风险区域表格
    highrisk_data = [
        ["监测区域", "高风险次数", "最高BI", "风险特征"],
        ["鹿寨县", "97", "43.59", "高风险频次最高"],
        ["融水县", "65", "60.00", "极端值较多"],
        ["柳城县", "60", "70.00", "最高达70"],
        ["柳南区", "47", "46.67", "城区中风险较高"],
        ["鱼峰区", "37", "56.52", "需重点关注"],
        ["融安县", "37", "200.00", "存在极端值"],
        ["城中区", "30", "32.14", "中心城区风险"]
    ]
    add_table_zh(doc, highrisk_data[1:], highrisk_data[0])
    
    doc.add_paragraph()
    
    add_heading_zh(doc, '3.3 风险等级分布', level=2)
    
    add_paragraph_zh(doc, 
        "从风险等级分布来看，72.8%的监测记录处于安全水平（BI<5），12.1%处于风险水平（5≤BI<10），15.1%处于高危水平（BI≥10）。虽然大部分监测点处于安全水平，但超过四分之一的监测记录存在登革热传播风险，需要引起重视。")
    
    # 插入风险分布图
    try:
        doc.add_picture(f'{output_dir}chart_04_risk_distribution.png', width=Inches(4.5))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_paragraph_zh(doc, "图4 布雷图指数风险等级分布", font_size=10, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    except:
        add_paragraph_zh(doc, "[图表：风险分布图]")
    
    # 插入热力图
    try:
        doc.add_picture(f'{output_dir}chart_05_heatmap.png', width=Inches(5.5))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_paragraph_zh(doc, "图5 布雷图指数年度-月份热力分布", font_size=10, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    except:
        add_paragraph_zh(doc, "[图表：热力分布图]")
    
    doc.add_page_break()
    
    # ========== 四、主要发现 ==========
    add_heading_zh(doc, '四、主要发现', level=1)
    
    findings = [
        "时间趋势方面：2021-2025年布雷图指数总体呈上升趋势，2025年首次超过风险阈值（BI=5），提示登革热传播风险显著增加。",
        "季节特征方面：夏季（6-8月）是布雷图指数高峰期，8月份平均BI达6.93，为全年最高；冬季布雷图指数最低，风险可控。",
        "空间分布方面：柳城县、鹿寨县为高风险区域，平均BI超过5；新建城区（柳东新区、阳和新区）风险较低。",
        "极端值方面：监测期间发现多例极端高值（BI>50），最高达200，提示局部地区存在严重的伊蚊孳生地。",
        "风险占比方面：27.2%的监测记录存在传播风险（BI≥5），其中15.1%为高危水平，需要采取紧急控制措施。",
        "区县差异方面：各县风险水平普遍高于市辖区，农村地区布雷图指数控制面临更大挑战。"
    ]
    
    for i, finding in enumerate(findings, 1):
        add_paragraph_zh(doc, f"{i}. {finding}", first_line_indent=0)
    
    doc.add_paragraph()
    
    # ========== 五、防控建议 ==========
    add_heading_zh(doc, '五、防控建议', level=1)
    
    add_heading_zh(doc, '5.1 重点时期防控', level=2)
    
    period_suggestions = [
        "每年6-9月为防控关键时期，需在5月底前完成全面的积水容器清理工作。",
        "8月份风险最高，应加强监测频次，从每月监测调整为每半月监测。",
        "台风、暴雨后要及时开展孳生地清除，防止积水容器增加。"
    ]
    
    for item in period_suggestions:
        add_paragraph_zh(doc, f"• {item}", first_line_indent=0)
    
    add_heading_zh(doc, '5.2 重点区域防控', level=2)
    
    area_suggestions = [
        "柳城县、鹿寨县应作为重点防控区域，建立常态化灭蚊机制。",
        "农村地区要加强健康教育，指导农户妥善管理积水容器。",
        "老旧小区、城中村等重点场所要定期开展翻盆倒罐行动。",
        "医院、学校、公园等公共场所要加强环境卫生管理。"
    ]
    
    for item in area_suggestions:
        add_paragraph_zh(doc, f"• {item}", first_line_indent=0)
    
    add_heading_zh(doc, '5.3 综合防控措施', level=2)
    
    measures = [
        "环境治理：清除各类积水容器，包括废弃轮胎、瓶罐、花盆托盘等，消除伊蚊孳生地。",
        "化学防制：布雷图指数≥10的区域应及时开展化学灭蚊，降低成蚊密度。",
        "生物防制：景观水体可放养食蚊鱼，减少蚊虫孳生。",
        "健康教育：开展登革热防控知识宣传，提高群众防蚊意识和参与度。",
        "监测预警：完善布雷图指数监测网络，及时发布风险预警信息。",
        "应急处置：建立登革热应急处置队伍，确保疫情发生时快速响应。"
    ]
    
    for item in measures:
        add_paragraph_zh(doc, f"• {item}", first_line_indent=0)
    
    doc.add_paragraph()
    
    # ========== 结语 ==========
    add_paragraph_zh(doc, 
        "综上所述，柳州市2021-2025年伊蚊布雷图指数监测数据显示，登革热传播风险总体可控但呈上升趋势，特别是2025年已超过风险阈值。建议各相关部门高度重视，采取综合防控措施，重点加强夏季高风险时期的监测和处置工作，切实保障人民群众健康安全。")
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # 添加报告结束标识
    end_para = doc.add_paragraph()
    end_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = end_para.add_run('—— 报告完 ——')
    set_chinese_font(run, font_name='SimSun', font_size=10)
    
    # 保存文档
    output_path = f'{output_dir}柳州市伊蚊布雷图监测分析报告.docx'
    doc.save(output_path)
    print(f"✓ Word报告已生成: {output_path}")
    
    return output_path

if __name__ == '__main__':
    generate_report()
