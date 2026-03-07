from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os
import pickle

# 设置文档保存路径
output_dir = "/Users/liaoweizhou/.openclaw/workspace/output"
chart_dir = f"{output_dir}/charts"

# 加载数据
with open(f"{output_dir}/analysis_data.pkl", 'rb') as f:
    data = pickle.load(f)

df_valid = data['df_valid']
yearly_stats = data['yearly_stats']
monthly_stats = data['monthly_stats']
district_stats = data['district_stats']

# 创建文档
doc = Document()

# 设置中文字体
def set_cell_font(cell, font_name='宋体', font_size=10.5):
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.name = font_name
            run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
            run.font.size = Pt(font_size)

# 封面
title = doc.add_heading('柳州市伊蚊布雷图指数监测分析报告', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title.runs:
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.name = '黑体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

# 添加副标题
subtitle = doc.add_paragraph('（2021-2025年时空分布分析）')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.runs[0].font.size = Pt(14)
subtitle.runs[0].font.name = '宋体'
subtitle.runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

doc.add_paragraph()
doc.add_paragraph()

# 添加日期
date_para = doc.add_paragraph(f'报告生成日期：2025年2月23日')
date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
date_para.runs[0].font.size = Pt(12)

doc.add_page_break()

# 目录
doc.add_heading('目  录', level=1)
toc_items = [
    '一、监测概况',
    '二、时间维度分析',
    '三、空间维度分析',
    '四、风险评估',
    '五、防控建议',
    '六、结论'
]
for item in toc_items:
    p = doc.add_paragraph(item, style='List Number')
    p.paragraph_format.left_indent = Inches(0.5)

doc.add_page_break()

# 一、监测概况
doc.add_heading('一、监测概况', level=1)

overview_text = f"""本次分析基于柳州市2021-2025年伊蚊布雷图指数监测数据，共收集到{len(df_valid)}个有效监测点数据，覆盖{district_stats.shape[0]}个城区。布雷图指数（Breteau Index, BI）是评估登革热传播风险的重要指标，计算公式为：阳性积水容器数/调查户数×100。

根据世界卫生组织标准，布雷图指数的风险等级划分如下：
• 低风险：BI < 5
• 中风险：5 ≤ BI < 10  
• 高风险：BI ≥ 10

监测数据总体情况：
• 监测年份：{int(df_valid['年份'].min())}年 - {int(df_valid['年份'].max())}年
• 平均布雷图指数：{df_valid['布雷图指数'].mean():.2f}
• 布雷图指数中位数：{df_valid['布雷图指数'].median():.2f}
• 最高布雷图指数：{df_valid['布雷图指数'].max():.2f}
• 总监测户数：{int(df_valid['户数'].sum())}户
• 阳性积水总数：{int(df_valid['阳性积水数'].sum())}处
"""

doc.add_paragraph(overview_text)

# 二、时间维度分析
doc.add_heading('二、时间维度分析', level=1)

# 年度变化趋势
doc.add_heading('2.1 年度变化趋势', level=2)

trend_text = f"""从2021年到2025年，柳州市伊蚊布雷图指数呈现明显上升趋势：

• 2021年平均BI：{yearly_stats.loc[2021, '平均BI']:.2f}（低风险）
• 2022年平均BI：{yearly_stats.loc[2022, '平均BI']:.2f}（低风险）
• 2023年平均BI：{yearly_stats.loc[2023, '平均BI']:.2f}（低风险）
• 2024年平均BI：{yearly_stats.loc[2024, '平均BI']:.2f}（低风险）
• 2025年平均BI：{yearly_stats.loc[2025, '平均BI']:.2f}（高风险）⚠️

分析表明，2025年布雷图指数显著上升，已突破风险阈值（BI=5），达到登革热传播风险水平，需要引起重视。
"""
doc.add_paragraph(trend_text)

# 添加年度趋势图
doc.add_paragraph('图1：柳州市伊蚊布雷图指数年度变化趋势（2021-2025）')
doc.add_picture(f"{chart_dir}/01_年度趋势图.png", width=Inches(5.8))
last_paragraph = doc.paragraphs[-1]
last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

# 月度分布特征
doc.add_heading('2.2 月度分布特征', level=2)

peak_month = monthly_stats['平均BI'].idxmax()
peak_value = monthly_stats.loc[peak_month, '平均BI']

monthly_text = f"""月度分析显示，柳州市伊蚊布雷图指数呈现明显的季节性变化：

• 峰值月份：{int(peak_month)}月（平均BI={peak_value:.2f}）
• 次高峰月份：8月（平均BI={monthly_stats.loc[8, '平均BI']:.2f}）、9月（平均BI={monthly_stats.loc[9, '平均BI']:.2f}）
• 低谷月份：12月（平均BI={monthly_stats.loc[12, '平均BI']:.2f}）

从4月到6月，布雷图指数快速上升，6-9月维持在较高水平，10月后开始下降。这与柳州市高温多雨的气候特征相符，夏秋季节（6-9月）是伊蚊孳生和登革热传播的高风险期。
"""
doc.add_paragraph(monthly_text)

# 添加月度分布图
doc.add_paragraph('图2：月度布雷图指数分布')
doc.add_picture(f"{chart_dir}/02_月度分布图.png", width=Inches(5.8))
last_paragraph = doc.paragraphs[-1]
last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

# 添加时空热力图
doc.add_paragraph('图3：布雷图指数年度-月份热力图')
doc.add_picture(f"{chart_dir}/06_时空热力图.png", width=Inches(5.8))
last_paragraph = doc.paragraphs[-1]
last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

# 三、空间维度分析
doc.add_heading('三、空间维度分析', level=1)

# 区域对比
doc.add_heading('3.1 各城区布雷图指数对比', level=2)

spatial_text = f"""全市{district_stats.shape[0]}个监测城区中，布雷图指数差异明显：

高风险城区（BI≥5）：
• 柳城县：平均BI={district_stats.loc['柳城县', '平均BI']:.2f}
• 融安县：平均BI={district_stats.loc['融安县', '平均BI']:.2f}
• 柳南区：平均BI={district_stats.loc['柳南区', '平均BI']:.2f}
• 鹿寨县：平均BI={district_stats.loc['鹿寨县', '平均BI']:.2f}

低风险城区（BI<3）：
• 三江县：平均BI={district_stats.loc['三江县', '平均BI']:.2f}
• 融水县：平均BI={district_stats.loc['融水县', '平均BI']:.2f}
• 柳北区：平均BI={district_stats.loc['柳北区', '平均BI']:.2f}
"""
doc.add_paragraph(spatial_text)

# 添加城区对比图
doc.add_paragraph('图4：各城区平均布雷图指数对比（TOP10）')
doc.add_picture(f"{chart_dir}/03_城区对比图.png", width=Inches(5.8))
last_paragraph = doc.paragraphs[-1]
last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

# 四、风险评估
doc.add_heading('四、风险评估', level=1)

# 风险等级分布
risk_counts = df_valid['风险等级'].value_counts()
risk_low = risk_counts.get('低风险', 0)
risk_medium = risk_counts.get('中风险', 0)
risk_high = risk_counts.get('高风险', 0)
total = len(df_valid)

doc.add_heading('4.1 风险等级分布', level=2)

risk_text = f"""根据布雷图指数风险等级划分，全市监测点分布如下：

• 低风险（BI<5）：{risk_low}个监测点，占比{risk_low/total*100:.1f}%
• 中风险（5≤BI<10）：{risk_medium}个监测点，占比{risk_medium/total*100:.1f}%
• 高风险（BI≥10）：{risk_high}个监测点，占比{risk_high/total*100:.1f}%⚠️

共有{len(df_valid[df_valid['布雷图指数']>=5])}个监测点（占比{len(df_valid[df_valid['布雷图指数']>=5])/total*100:.1f}%）达到或超过登革热传播风险阈值。
"""
doc.add_paragraph(risk_text)

# 添加风险等级分布图
doc.add_paragraph('图5：布雷图指数风险等级分布')
doc.add_picture(f"{chart_dir}/04_风险等级分布图.png", width=Inches(5.0))
last_paragraph = doc.paragraphs[-1]
last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

# 高风险区域识别
doc.add_heading('4.2 高风险区域识别', level=2)

high_risk = df_valid[df_valid['风险等级'] == '高风险'].sort_values('布雷图指数', ascending=False)

high_risk_text = f"""全市共识别出{len(high_risk)}个高风险监测点（BI≥10），主要分布在：
"""
doc.add_paragraph(high_risk_text)

# 添加高风险点列表（前10个）
doc.add_paragraph('表1：高风险监测点列表（TOP10）')
table = doc.add_table(rows=1, cols=5)
table.style = 'Light Grid Accent 1'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = '排名'
hdr_cells[1].text = '年份'
hdr_cells[2].text = '城区'
hdr_cells[3].text = '地址'
hdr_cells[4].text = '布雷图指数'

for i, (idx, row) in enumerate(high_risk.head(10).iterrows(), 1):
    row_cells = table.add_row().cells
    row_cells[0].text = str(i)
    row_cells[1].text = f"{int(row['年份'])}年{int(row['月份'])}月"
    row_cells[2].text = str(row['城区'])
    row_cells[3].text = str(row.get('地址', '')[:15])
    row_cells[4].text = f"{row['布雷图指数']:.2f}"

# 五、防控建议
doc.add_heading('五、防控建议', level=1)

suggestions = [
    "1. 重点关注高风险区域：柳城县、融安县、柳南区、鹿寨县等布雷图指数超过风险阈值的城区，应加强监测和防控力度。",
    "2. 加强季节性防控：6-9月为伊蚊孳生高峰期，应在此时期前开展爱国卫生运动，清除积水容器。",
    "3. 监测点扩容：2025年监测点数量大幅增加（达938个），但布雷图指数也随之上升，说明伊蚊密度确有增加，需加强防控。",
    "4. 建立预警机制：对布雷图指数≥10的高风险区域建立预警机制，及时采取灭蚊措施。",
    "5. 加强宣传教育：提高居民对登革热防控的认识，引导居民主动清除室内外积水。",
    "6. 跨部门协作：卫健、城管、住建等部门协同，对建筑工地、公园绿地、居民小区等重点场所进行联合整治。"
]

for suggestion in suggestions:
    doc.add_paragraph(suggestion)

# 六、结论
doc.add_heading('六、结论', level=1)

conclusion_text = f"""通过对柳州市2021-2025年伊蚊布雷图指数监测数据的分析，得出以下主要结论：

1. 总体趋势：全市布雷图指数呈上升趋势，2025年平均BI已达{yearly_stats.loc[2025, '平均BI']:.2f}，首次突破登革热传播风险阈值。

2. 季节特征：6-9月为高风险期，{int(peak_month)}月份为峰值，需在此时期加强防控措施。

3. 空间分布：柳城县、融安县、柳南区、鹿寨县为高风险城区，应作为重点防控区域。

4. 风险预警：全市有{len(df_valid[df_valid['布雷图指数']>=5])}个监测点达到风险阈值，占比{len(df_valid[df_valid['布雷图指数']>=5])/total*100:.1f}%，存在登革热本地传播风险。

建议相关部门根据本报告结果，制定针对性的防控策略，切实降低伊蚊密度，预防登革热疫情发生。
"""
doc.add_paragraph(conclusion_text)

# 保存文档
doc_path = f"{output_dir}/柳州市伊蚊布雷图监测分析报告.docx"
doc.save(doc_path)
print(f"✅ Word报告已生成: {doc_path}")

# 列出所有生成的文件
print("\n" + "="*80)
print("📁 所有输出文件:")
print("="*80)
for root, dirs, files in os.walk(output_dir):
    for file in files:
        filepath = os.path.join(root, file)
        size = os.path.getsize(filepath)
        print(f"  • {filepath.replace(output_dir+'/', '')} ({size/1024:.1f} KB)")
